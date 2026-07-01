"""Render generated SWMS data into the HRCW Word template."""

from __future__ import annotations

import io
import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from app.exceptions import DocumentRenderError
from app.schemas.render import RenderSwmsRequest

DEFAULT_TEMPLATE_PATH = (
    Path(__file__).resolve().parents[2] / "templates" / "swms-hrcw-risk-columns.docx"
)

HEADER_TABLE_INDEX = 0
STEPS_TABLE_INDEX = 4
STEPS_DATA_START_ROW = 3

STEP_COLUMNS = {
    "step_no": 0,
    "job_task": 1,
    "hazard": 2,
    "risk_level": 3,
    "risk_control": 4,
    "post_risk": 6,
    "responsible": 7,
}

RISK_CONTROL_COLUMN = STEP_COLUMNS["risk_control"]
HIERARCHY_COLUMN = 5

COMBINED_RISK_CONTROL_HEADER = (
    "RISK CONTROL\n"
    "How will the risk be controlled effectively?\n\n"
    "Hierarchy of Risk Control (informative):\n"
    "1. Eliminate the hazard\n"
    "2. Substitute the hazard\n"
    "3. Isolate the hazard from people\n"
    "4. Engineering controls\n"
    "5. Administrative controls\n"
    "6. Personal protective equipment (PPE)"
)

HEADER_CELL_MAP: dict[tuple[int, int], str] = {
    (1, 2): "companyName",
    (5, 2): "siteLocation",
    (7, 2): "jobActivitiesText",
    (9, 2): "plantEquipmentText",
    (7, 17): "swmsIssueDate",
}


def resolve_template_path(template_path: str | None = None) -> Path:
    """Resolves the SWMS Word template path from argument or environment."""
    if template_path:
        path = Path(template_path)
        if not path.is_file():
            raise DocumentRenderError(f"SWMS template not found: {path}")
        return path

    env_path = Path(__file__).resolve().parents[2] / "templates" / "swms-hrcw-risk-columns.docx"
    if env_path.is_file():
        return env_path

    if DEFAULT_TEMPLATE_PATH.is_file():
        return DEFAULT_TEMPLATE_PATH

    raise DocumentRenderError("SWMS Word template is not configured")


def _format_controls(controls: list[str]) -> str:
    if not controls:
        return ""
    return "\n".join(f"• {control.strip()}" for control in controls if control.strip())


def _format_combined_risk_control(controls: list[str]) -> str:
    """Combines actionable controls with informative hierarchy guidance in one cell."""
    controls_text = _format_controls(controls)
    hierarchy = _infer_hierarchy_control(controls)
    if not hierarchy:
        return controls_text

    informative = f"Hierarchy (informative): {hierarchy}"
    if controls_text:
        return f"{controls_text}\n\n{informative}"
    return informative


def _infer_hierarchy_control(controls: list[str]) -> str:
    text = " ".join(controls).lower()
    if re.search(r"\b(eliminat|remove hazard|do not perform)\b", text):
        return "1. Eliminate the hazard"
    if re.search(r"\b(substitut|replace with)\b", text):
        return "2. Substitute the hazard"
    if re.search(r"\b(isolate|barricad|exclusion zone|lockout|tagout|lock out|tag out)\b", text):
        return "3. Isolate the hazard from people"
    if re.search(r"\b(engineer|guard|ventilat|interlock|machine guard)\b", text):
        return "4. Engineering controls"
    if re.search(r"\b(procedure|training|supervis|induction|permit|administrative)\b", text):
        return "5. Administrative controls"
    if re.search(r"\b(ppe|glove|helmet|harness|boot|goggle|mask|respirator|hi[\s-]?vis)\b", text):
        return "6. Personal protective equipment (PPE)"
    return ""


def _set_vertical_merge(cell, *, restart: bool) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    for child in tc_pr.findall(qn("w:vMerge")):
        tc_pr.remove(child)

    v_merge = OxmlElement("w:vMerge")
    if restart:
        v_merge.set(qn("w:val"), "restart")
    tc_pr.append(v_merge)


def _set_cell_text_top(cell, text: str) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    cell.text = text
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _clear_cell(cell) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    cell.text = ""


def _merge_column_vertically(table, column_index: int, start_row: int, end_row: int) -> None:
    if start_row >= end_row:
        return

    top_cell = table.rows[start_row].cells[column_index]
    _set_vertical_merge(top_cell, restart=True)
    top_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    for row_index in range(start_row + 1, end_row + 1):
        continuation_cell = table.rows[row_index].cells[column_index]
        _set_vertical_merge(continuation_cell, restart=False)
        _clear_cell(continuation_cell)


def _merge_cells_horizontally(row, start_column: int, end_column: int) -> None:
    if start_column >= end_column:
        return

    start_cell = row.cells[start_column]
    for column_index in range(start_column + 1, end_column + 1):
        start_cell.merge(row.cells[column_index])


def _prepare_steps_table_header(table) -> None:
    """Merges risk-control and hierarchy header cells into one informative column."""
    header_row = table.rows[2]
    _merge_cells_horizontally(header_row, RISK_CONTROL_COLUMN, HIERARCHY_COLUMN)
    _set_cell_text_top(header_row.cells[RISK_CONTROL_COLUMN], COMBINED_RISK_CONTROL_HEADER)


def _build_header_context(request: RenderSwmsRequest) -> dict[str, str]:
    job_lines = [request.title.strip(), *request.jobActivities]
    job_text = "\n".join(line.strip() for line in job_lines if line.strip())
    plant_text = "\n".join(item.strip() for item in request.plantEquipment if item.strip())
    issue_date = request.swmsIssueDate or date.today().strftime("%d %B %Y")

    return {
        "companyName": (request.companyName or "").strip(),
        "siteLocation": request.siteLocation.strip(),
        "jobActivitiesText": job_text,
        "plantEquipmentText": plant_text,
        "swmsIssueDate": issue_date,
    }


def _fill_header_table(document: Document, request: RenderSwmsRequest) -> None:
    if len(document.tables) <= HEADER_TABLE_INDEX:
        raise DocumentRenderError("Header table not found in SWMS template")

    header_table = document.tables[HEADER_TABLE_INDEX]
    context = _build_header_context(request)

    for (row_index, column_index), field_name in HEADER_CELL_MAP.items():
        if row_index >= len(header_table.rows):
            continue
        row = header_table.rows[row_index]
        if column_index >= len(row.cells):
            continue
        _set_cell_text_top(row.cells[column_index], context[field_name])

    if document.paragraphs:
        title_text = request.title.strip()
        if title_text:
            document.paragraphs[0].text = f"SAFE WORK METHOD STATEMENT (SWMS)\n{title_text}"


def _remove_steps_data_rows(table) -> None:
    while len(table.rows) > STEPS_DATA_START_ROW:
        table._tbl.remove(table.rows[-1]._tr)


def _fill_steps_table(document: Document, request: RenderSwmsRequest) -> None:
    if len(document.tables) <= STEPS_TABLE_INDEX:
        raise DocumentRenderError("SWMS steps table not found in template")

    table = document.tables[STEPS_TABLE_INDEX]
    _prepare_steps_table_header(table)
    _remove_steps_data_rows(table)

    for step in request.steps:
        hazard_count = len(step.hazards)
        start_row_index = len(table.rows)

        for hazard_index, hazard in enumerate(step.hazards):
            row = table.add_row()
            cells = row.cells

            if hazard_index == 0:
                _set_cell_text_top(cells[STEP_COLUMNS["step_no"]], step.stepNo)
                _set_cell_text_top(cells[STEP_COLUMNS["job_task"]], step.jobTaskElement)
            else:
                _clear_cell(cells[STEP_COLUMNS["step_no"]])
                _clear_cell(cells[STEP_COLUMNS["job_task"]])

            _set_cell_text_top(cells[STEP_COLUMNS["hazard"]], hazard.hazard)
            _set_cell_text_top(cells[STEP_COLUMNS["risk_level"]], hazard.riskLevel)
            _set_cell_text_top(
                cells[STEP_COLUMNS["risk_control"]],
                _format_combined_risk_control(hazard.controls),
            )
            _merge_cells_horizontally(row, RISK_CONTROL_COLUMN, HIERARCHY_COLUMN)
            _set_cell_text_top(cells[STEP_COLUMNS["post_risk"]], hazard.postRiskLevel)
            _set_cell_text_top(cells[STEP_COLUMNS["responsible"]], hazard.responsiblePerson)

            if len(cells) > 8:
                _set_cell_text_top(cells[8], hazard.responsiblePerson)

        end_row_index = start_row_index + hazard_count - 1
        if hazard_count > 1:
            _merge_column_vertically(table, STEP_COLUMNS["step_no"], start_row_index, end_row_index)
            _merge_column_vertically(table, STEP_COLUMNS["job_task"], start_row_index, end_row_index)


def render_swms_document_bytes(
    request: RenderSwmsRequest,
    template_path: str | None = None,
) -> bytes:
    """
    Renders a SWMS Word document from the template and structured step data.

    Args:
        request: Header metadata and generated steps/hazards.
        template_path: Optional override for the template file path.

    Returns:
        DOCX file bytes.
    """
    path = resolve_template_path(template_path)
    document = Document(str(path))

    _fill_header_table(document, request)
    _fill_steps_table(document, request)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def render_swms_document_bytes_from_dict(
    payload: dict,
    template_path: str | None = None,
) -> bytes:
    """Validates payload and renders a SWMS Word document."""
    request = RenderSwmsRequest.model_validate(payload)
    return render_swms_document_bytes(request, template_path=template_path)
