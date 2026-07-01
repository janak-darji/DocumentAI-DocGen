"""Tests for SWMS Word rendering."""

from __future__ import annotations

import io

from docx import Document
from docx.oxml.ns import qn

from app.schemas.render import RenderSwmsRequest
from app.services.swms_renderer import render_swms_document_bytes, resolve_template_path


def _sample_request() -> RenderSwmsRequest:
    return RenderSwmsRequest(
        title="Switchboard replacement",
        siteLocation="123 George St, Brisbane",
        jobActivities=["Isolate supply", "Install switchboard"],
        plantEquipment=["Multimeter", "Insulated tools"],
        tradeName="Electrician",
        activityType="Not Applicable",
        steps=[
            {
                "stepNo": "1",
                "jobTaskElement": "Isolate and lock out electrical supply",
                "sequencePosition": 1,
                "hazards": [
                    {
                        "hazard": "Electric shock",
                        "riskLevel": "1",
                        "controls": ["Lockout/tagout", "Test for dead"],
                        "postRiskLevel": "3",
                        "responsiblePerson": "Licensed electrician",
                    },
                    {
                        "hazard": "Arc flash",
                        "riskLevel": "1",
                        "controls": ["Wear arc-rated PPE"],
                        "postRiskLevel": "2",
                        "responsiblePerson": "Licensed electrician",
                    },
                ],
            },
            {
                "stepNo": "2",
                "jobTaskElement": "Install new switchboard",
                "sequencePosition": 2,
                "hazards": [
                    {
                        "hazard": "Manual handling injury",
                        "riskLevel": "2",
                        "controls": ["Use mechanical aids", "Team lift"],
                        "postRiskLevel": "3",
                        "responsiblePerson": "Leading hand",
                    }
                ],
            },
        ],
    )


def test_resolve_template_path_finds_bundled_template() -> None:
    path = resolve_template_path()
    assert path.name == "swms-hrcw-risk-columns.docx"


def test_render_swms_document_bytes_produces_steps_with_vertical_merge() -> None:
    rendered = render_swms_document_bytes(_sample_request())
    assert rendered.startswith(b"PK")

    document = Document(io.BytesIO(rendered))
    steps_table = document.tables[4]

    assert len(steps_table.rows) == 6  # 3 header rows + 3 hazard rows

    first_step_top = steps_table.rows[3].cells[0].text
    first_step_continuation_tc = steps_table.rows[4]._tr.findall(qn("w:tc"))[0]
    first_step_continuation_vmerge = first_step_continuation_tc.xpath("./w:tcPr/w:vMerge")
    assert first_step_top == "1"
    assert len(first_step_continuation_vmerge) == 1

    assert "Electric shock" in steps_table.rows[3].cells[2].text
    assert "Arc flash" in steps_table.rows[4].cells[2].text
    assert steps_table.rows[3].cells[1].text == "Isolate and lock out electrical supply"
    second_task_tc = steps_table.rows[4]._tr.findall(qn("w:tc"))[1]
    assert second_task_tc.xpath("./w:tcPr/w:vMerge")

    assert steps_table.rows[5].cells[0].text == "2"
    assert "Manual handling injury" in steps_table.rows[5].cells[2].text

    risk_control_tc = steps_table.rows[3]._tr.findall(qn("w:tc"))[4]
    grid_span = risk_control_tc.xpath("./w:tcPr/w:gridSpan")
    assert grid_span and grid_span[0].get(qn("w:val")) == "2"
    assert "Lockout/tagout" in steps_table.rows[3].cells[4].text
    assert "Hierarchy (informative):" in steps_table.rows[3].cells[4].text


def test_render_swms_document_bytes_fills_header_fields() -> None:
    rendered = render_swms_document_bytes(_sample_request())
    document = Document(io.BytesIO(rendered))
    header_table = document.tables[0]

    assert "123 George St, Brisbane" in header_table.rows[5].cells[2].text
    assert "Switchboard replacement" in header_table.rows[7].cells[2].text
    assert "Multimeter" in header_table.rows[9].cells[2].text
